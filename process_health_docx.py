import openai
import os
import json
from docx import Document
from dotenv import load_dotenv


# --- Set your OpenAI API key ---
# openai.api_key = os.getenv("OPENAI_API_KEY")  # or paste directly: "sk-..."
# Load environment variables from .env
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

# --- Load the .docx file and extract plain text ---
def load_docx_text(file_path):
    doc = Document(file_path)
    lines = []
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                lines.append(" | ".join(row_text))
    return "\n".join(lines)

# --- Prepare the prompt for GPT ---
def build_prompt(text):
    return f"""
You are a medical document analysis assistant.

Given the following health checkup document text:

\"\"\"{text}\"\"\\"

Extract all meaningful medical key-value pairs, such as lab test names and their corresponding values, and return them in flat JSON format like:

{{
  "LDL Cholesterol": "84.3 mg/dL (< 135.3 mg/dL)",
  "Total Cholesterol": "185.5 mg/dL (< 201.1 mg/dL)",
  "HDL Cholesterol": "76.6 mg/dL (≥ 38.7 mg/dL)",
  "Non-HDL Cholesterol": "108.7 mg/dL (< 162.4 mg/dL)",
  "Triglycerides": "144.4 mg/dL (< 150.6 mg/dL)",
  "A1c": "5.4% (< 6.0%)",
  "eGFR": "72 (≥ 60)",
  "Urea": "49.2 mg/dL (15.0 – 48.6 mg/dL)",
  "Iron": "12 (7 – 29)",
  "Vitamin B12": "149 (148–220)",
  "PSA": "1.68 (< 3.5)",
  "DHEAS": "4.1 (1.9 – 8.4)",
  "WBC Count": "4.1 x10⁹/L (4.5 – 11.0 x10⁹/L)",
  "RBC Count": "4.9 x10¹²/L (4.4 – 5.9 x10¹²/L)",
  "Hemoglobin": "152 g/L (140 – 180 g/L)",
  "Lymphocytes": "0.8 x10⁹/L (1.0 – 3.3 x10⁹/L)",
  "Ferritin": "473 (> 220)",
  "Platelets": "170 x10⁹/L (140 – 440 x10⁹/L)"
}}

For each test, combine value, unit, and reference range into a single string in the format: "<value> <unit> (<reference range>)". If reference range is not available, omit the parentheses.

When possible, convert lab values and reference ranges to U.S. standard units using the following conversion factors:
- Cholesterol (mmol/L to mg/dL): multiply by 38.67
- Triglycerides (mmol/L to mg/dL): multiply by 88.57
- Creatinine (µmol/L to mg/dL): divide by 88.4
- Uric Acid (µmol/L to mg/dL): divide by 59.48
- Urea (mmol/L to mg/dL): multiply by 6.0
- Glucose (mmol/L to mg/dL): multiply by 18.0

If a test is not listed above, keep the original unit. Always show the final value and reference range in the same format as the example above, and round converted values to one decimal place.
If the test name is in Chinese, convert the value and reference range to U.S. standard units and output the result in the same format, using mg/dL, ng/mL, etc. for the relevant analytes.
Do not group by section. Keep test names and values in original language (e.g., Chinese or English).
"""

# --- Call OpenAI GPT to extract and convert data ---
def extract_medical_data(doc_text, model="gpt-4"):
    prompt = build_prompt(doc_text)
    client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0
    )
    content = response.choices[0].message.content
    try:
        data = json.loads(content)
        # If the data is nested, flatten it
        def flatten(d):
            flat = {}
            for k, v in d.items():
                if isinstance(v, dict):
                    value = v.get("value", "")
                    unit = v.get("unit", "")
                    ref = v.get("reference_range", "")
                    s = f"{value} {unit}".strip()
                    if ref:
                        s += f" ({ref})"
                    flat[k] = s.strip()
                else:
                    flat[k] = v
            return flat
        # Only flatten if needed
        if any(isinstance(v, dict) for v in data.values()):
            data = flatten(data)
        # Post-process for Chinese analytes
        data = convert_chinese_lab_units(data)
        return data
    except Exception:
        return content  # fallback to raw string if not pure JSON

def convert_chinese_lab_units(result_dict):
    # Conversion factors
    conversion_map = {
        "高密度脂蛋白胆固醇": (38.67, "mg/dL"),
        "低密度脂蛋臼胆固醇": (38.67, "mg/dL"),
        "总胆固酪": (38.67, "mg/dL"),
        "甘油三醋": (88.57, "mg/dL"),
        "尿素": (6.0, "mg/dL"),
        "尿酸": (59.48, "mg/dL"),
        "肌酐": (88.4, "mg/dL"),
    }
    import re
    def convert_value(val, factor):
        # Find float in string
        m = re.search(r"([\d.]+)", val)
        if m:
            return round(float(m.group(1)) * factor, 1)
        return val
    def convert_range(rng, factor):
        # Convert ranges like "3.10-8.80 mmol/L" or ">1.04 mmol/L" or "<5.18 mmol/L"
        # Replace all numbers
        def repl(m):
            return str(round(float(m.group(0)) * factor, 1))
        return re.sub(r"[\d.]+", repl, rng)
    for k, v in result_dict.items():
        if k in conversion_map:
            factor, new_unit = conversion_map[k]
            # Parse value and reference range
            m = re.match(r"([\d.]+)[^\d]*([a-zA-Z/]+)? ?\(([^)]*)\)", v)
            if m:
                value = convert_value(m.group(1), factor)
                ref = convert_range(m.group(3), factor)
                result_dict[k] = f"{value} {new_unit} ({ref})"
            else:
                # Try to parse value only
                m2 = re.match(r"([\d.]+)[^\d]*([a-zA-Z/]+)?", v)
                if m2:
                    value = convert_value(m2.group(1), factor)
                    result_dict[k] = f"{value} {new_unit}"
    # Special cases for µmol/L to mg/dL
    for k in ["尿酸", "肌酐"]:
        if k in result_dict:
            v = result_dict[k]
            factor = conversion_map[k][0]
            new_unit = conversion_map[k][1]
            m = re.match(r"([\d.]+)[^\d]*([a-zA-Z/µ]+)? ?\(([^)]*)\)", v)
            if m:
                value = round(float(m.group(1)) / factor, 1)
                ref = re.sub(r"[\d.]+", lambda m: str(round(float(m.group(0)) / factor, 1)), m.group(3))
                result_dict[k] = f"{value} {new_unit} ({ref})"
    return result_dict

# --- Save result to a file ---
def save_result(result, output_path="output.json"):
    with open(output_path, "w", encoding="utf-8") as f:
        if isinstance(result, dict):
            json.dump(result, f, ensure_ascii=False, indent=2)
        else:
            f.write(result)

# --- Main execution ---
if __name__ == "__main__":
    #FILE_PATH = "血脂.docx"  # Replace with your file path
    FILE_PATH = "CanadaBloodWork.docx"  # Replace with your file path
    OUTPUT_PATH = "parsed_medical_data.json"

    text = load_docx_text(FILE_PATH)
    result = extract_medical_data(text)
    if result:
        print(json.dumps(result, indent=2, ensure_ascii=False))
    # save_result(result, OUTPUT_PATH)

    # print(f"✅ Extracted data saved to: {OUTPUT_PATH}")
